from openpyxl import load_workbook
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

path='/mnt/data/lucro_receita_despesa(1).xlsx'

# Load data
resultado = pd.read_excel(path, sheet_name='Resultado')
receitas = pd.read_excel(path, sheet_name='Receitas')
despesas = pd.read_excel(path, sheet_name='Despesas')

# Prepare dates
resultado['mês'] = pd.to_datetime(resultado['mês'])
resultado['mês_str'] = resultado['mês'].dt.strftime('%b/%Y')

# KPIs
total_receita = resultado['receita'].sum()
total_despesa = resultado['despesa'].sum()
total_lucro = resultado['lucro'].sum()
margem_media = (resultado['lucro'].sum() / resultado['receita'].sum()) * 100
crescimento_receita = ((resultado['receita'].iloc[-1] - resultado['receita'].iloc[0]) / resultado['receita'].iloc[0]) * 100

# Chart 1: Receita x Despesa x Lucro
plt.figure(figsize=(8,4))
plt.plot(resultado['mês_str'], resultado['receita'], marker='o')
plt.plot(resultado['mês_str'], resultado['despesa'], marker='o')
plt.plot(resultado['mês_str'], resultado['lucro'], marker='o')
plt.xlabel('Mês')
plt.ylabel('Valor')
plt.title('Receita x Despesa x Lucro')
chart1 = '/mnt/data/chart_resultado.png'
plt.savefig(chart1, bbox_inches='tight')
plt.close()

# Receita por produto
receitas_melt = receitas.melt(id_vars='Produto', var_name='Mes', value_name='Receita')
receitas_melt['Mes'] = pd.to_datetime(receitas_melt['Mes'])

produto_total = receitas_melt.groupby('Produto')['Receita'].sum().sort_values(ascending=False)

plt.figure(figsize=(6,4))
produto_total.plot(kind='bar')
plt.ylabel('Receita Total')
plt.title('Receita Total por Produto')
chart2 = '/mnt/data/chart_produtos.png'
plt.savefig(chart2, bbox_inches='tight')
plt.close()

# Despesas por categoria
desp_cols = [c for c in despesas.columns if '2026' in str(c)]
despesas['Total'] = despesas[desp_cols].sum(axis=1)

plt.figure(figsize=(7,4))
plt.bar(despesas['categoria'], despesas['Total'])
plt.ylabel('Despesa Total')
plt.title('Despesas por Categoria')
chart3 = '/mnt/data/chart_despesas.png'
plt.savefig(chart3, bbox_inches='tight')
plt.close()

# Create report
doc = Document()

title = doc.add_heading('Análise Financeira Completa', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Resumo Executivo', level=2)
doc.add_paragraph(
    f'O período analisado apresentou receita total de R$ {total_receita:,.0f}, '
    f'despesas totais de R$ {total_despesa:,.0f} e lucro acumulado de R$ {total_lucro:,.0f}. '
    f'A margem média de lucro foi de {margem_media:.2f}%.'
)

doc.add_heading('2. Indicadores Financeiros', level=2)

table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Indicador'
hdr_cells[1].text = 'Valor'

indicators = [
    ('Receita Total', f'R$ {total_receita:,.0f}'),
    ('Despesa Total', f'R$ {total_despesa:,.0f}'),
    ('Lucro Total', f'R$ {total_lucro:,.0f}'),
    ('Margem Média', f'{margem_media:.2f}%'),
    ('Crescimento da Receita', f'{crescimento_receita:.2f}%')
]

for ind, val in indicators:
    row_cells = table.add_row().cells
    row_cells[0].text = ind
    row_cells[1].text = val

doc.add_heading('3. Evolução Financeira', level=2)
doc.add_paragraph(
    'A análise temporal mostra crescimento da receita ao longo do período, '
    'com oscilações controladas nas despesas e manutenção do lucro positivo.'
)
doc.add_picture(chart1, width=Inches(6.5))

doc.add_heading('4. Análise de Receitas', level=2)
top_produto = produto_total.idxmax()
doc.add_paragraph(
    f'O produto com maior contribuição de receita foi o Produto {top_produto}. '
    'A distribuição entre produtos demonstra relativa diversificação de receitas.'
)
doc.add_picture(chart2, width=Inches(5.8))

doc.add_heading('5. Análise de Despesas', level=2)
maior_categoria = despesas.sort_values('Total', ascending=False).iloc[0]['categoria']
doc.add_paragraph(
    f'A categoria com maior volume de despesas foi {maior_categoria}. '
    'Observa-se necessidade de monitoramento principalmente nas categorias '
    'com crescimento acelerado.'
)
doc.add_picture(chart3, width=Inches(6))

doc.add_heading('6. Insights Estratégicos', level=2)

insights = [
    'A empresa manteve lucro positivo em todos os meses analisados.',
    'O crescimento da receita foi superior ao crescimento das despesas no período.',
    'Existe concentração relevante de despesas em RH e Infraestrutura.',
    'Marketing apresentou aumento expressivo em março, sugerindo possível campanha ou expansão.',
    'A margem de lucro permanece saudável, acima de 20% no consolidado.'
]

for item in insights:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7. Recomendações', level=2)

recomendacoes = [
    'Monitorar o crescimento das despesas de Marketing e TI.',
    'Criar metas de margem por produto.',
    'Implementar forecast financeiro mensal.',
    'Avaliar indicadores de ROI por categoria de despesa.',
    'Expandir análise para fluxo de caixa e projeções futuras.'
]

for item in recomendacoes:
    doc.add_paragraph(item, style='List Number')

output_path = '/mnt/data/analise_financeira_completa.docx'
doc.save(output_path)

print(f'Relatório salvo em: {output_path}')

